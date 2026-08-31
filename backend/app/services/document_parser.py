from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class UnsupportedDocument(ValueError):
    pass


def parse_document(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise UnsupportedDocument("无法识别TXT编码")
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if not text:
            raise UnsupportedDocument("PDF未提取到文本，扫描件需要OCR")
        return text
    if suffix == ".docx":
        document = Document(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()).strip()
        if not text:
            raise UnsupportedDocument("DOCX未提取到文本")
        return text
    raise UnsupportedDocument(f"暂不支持 {suffix or '无扩展名'} 文件")
